"""Phase-1 preflight — template upload hardening (deferred Phase-0 item).

Office templates (xlsx/docx) are ZIP archives, so they carry zip-bomb / archive
abuse risks a plain image doesn't, and the old check trusted extension/kind. These
tests prove signature-mismatch rejection and archive-bomb caps (checked from the
archive's central directory, without decompressing).
"""
import io

import docx
import openpyxl
import pytest
from fastapi import HTTPException

from app import uploads

XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _xlsx() -> bytes:
    wb = openpyxl.Workbook()
    wb.active.append(["Name", "Amount", "Date"])
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()


def _docx() -> bytes:
    d = docx.Document(); d.add_paragraph("Name:"); d.add_paragraph("Total:")
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


_PDF = b"%PDF-1.4\n1 0 obj<<>>endobj\ntrailer<<>>\n%%EOF"
_HTML = b"<html><body><input name='vendor'></body></html>"


def test_valid_office_pdf_html_pass():
    uploads.validate_template_upload(_xlsx(), XLSX)      # no raise
    uploads.validate_template_upload(_docx(), DOCX)
    uploads.validate_template_upload(_PDF, "application/pdf")
    uploads.validate_template_upload(_HTML, "text/html")


def test_signature_mismatch_rejected():
    # A PDF uploaded as xlsx, and a ZIP uploaded as pdf → 415 (signature wins).
    with pytest.raises(HTTPException) as e1:
        uploads.validate_template_upload(_PDF, XLSX)
    assert e1.value.status_code == 415
    with pytest.raises(HTTPException) as e2:
        uploads.validate_template_upload(_xlsx(), "application/pdf")
    assert e2.value.status_code == 415


def test_html_must_be_text_not_disguised_binary():
    with pytest.raises(HTTPException) as e:
        uploads.validate_template_upload(_xlsx(), "text/html")   # a zip as html
    assert e.value.status_code == 415


def test_malformed_office_archive_rejected():
    fake = b"PK\x03\x04" + b"not a real zip central directory"
    with pytest.raises(HTTPException) as e:
        uploads.validate_template_upload(fake, XLSX)
    assert e.value.status_code == 400


def test_too_many_entries_rejected(monkeypatch):
    monkeypatch.setattr(uploads, "MAX_TEMPLATE_ENTRIES", 1)   # xlsx has several parts
    with pytest.raises(HTTPException) as e:
        uploads.validate_template_upload(_xlsx(), XLSX)
    assert e.value.status_code == 413


def test_uncompressed_size_cap_rejected(monkeypatch):
    monkeypatch.setattr(uploads, "MAX_TEMPLATE_UNCOMPRESSED", 100)  # bytes
    with pytest.raises(HTTPException) as e:
        uploads.validate_template_upload(_xlsx(), XLSX)
    assert e.value.status_code == 413


def test_compression_ratio_bomb_rejected(monkeypatch):
    # A highly-compressible entry (lots of zeros) trips a low ratio cap without
    # allocating anything large here.
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("payload.xml", b"0" * 500_000)   # compresses to ~KB
    data = buf.getvalue()
    monkeypatch.setattr(uploads, "MAX_TEMPLATE_RATIO", 5)
    monkeypatch.setattr(uploads, "MAX_TEMPLATE_UNCOMPRESSED", 10 * 1024 * 1024)
    with pytest.raises(HTTPException) as e:
        uploads.validate_template_upload(data, XLSX)
    assert e.value.status_code == 413
