"""Polished document reproduction — Path A.

Take an extracted result (any of our document_type shapes) and re-render it as
a clean, branded HTML/PDF/DOCX that customers can hand to a client. The
template lives in templates/repro/<doctype>.html; this module just normalizes
the extracted shape into a flat dict the template can read directly, then
renders.

Path A (this file): pre-built templates, looks consistent per doctype.
Path B (future):     ask the model to ALSO produce a layout HTML matching the
                     source image, then fill it with the extracted data.
                     Pure addition — the template fallback still works.
"""
from __future__ import annotations

from typing import Any
from pathlib import Path
from jinja2 import Environment, FileSystemLoader, select_autoescape

_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates" / "repro"
_jinja = Environment(
    loader=FileSystemLoader(str(_TEMPLATES_DIR)),
    autoescape=select_autoescape(["html"]),
    trim_blocks=True, lstrip_blocks=True,
)

# Document types we have a polished template for. Add to this list when a new
# templates/repro/<slug>.html is created.
SUPPORTED = {"invoice", "prescription", "business_card", "national_id", "bank_statement"}


def _val(node: Any, default: str = "") -> str:
    """Pull `value` out of a {value, confidence} cell OR return primitive."""
    if node is None:
        return default
    if isinstance(node, dict):
        # Prefer the normalised ISO date if we wrote one; fall through to raw value.
        if node.get("value_iso"):
            return str(node["value_iso"])
        if node.get("value") not in (None, ""):
            return str(node["value"])
        amt = node.get("amount")
        if isinstance(amt, dict) and amt.get("value") is not None:
            return str(amt["value"])
        return default
    return str(node)


def _amt(node: Any, default: str = "") -> str:
    """Pull the human-readable amount: prefer normalized {value, currency}."""
    if isinstance(node, dict):
        amt = node.get("amount")
        if isinstance(amt, dict) and amt.get("value") is not None:
            cur = amt.get("currency") or ""
            return f"{amt['value']:,.2f} {cur}".strip()
        v = node.get("value")
        return str(v) if v is not None else default
    return str(node) if node is not None else default


def _flatten_invoice(data: dict) -> dict:
    """Normalize the many shapes we emit (rich object, flat fields, mixed
    layout) into a single dict the invoice template can render against."""
    if not isinstance(data, dict):
        return {}
    header = data.get("header") if isinstance(data.get("header"), dict) else {}
    totals = data.get("totals") if isinstance(data.get("totals"), dict) else {}
    items_blk = data.get("line_items") or {}
    supplier = data.get("supplier") if isinstance(data.get("supplier"), dict) else {}
    customer = data.get("customer") if isinstance(data.get("customer"), dict) else {}
    fields = data.get("fields") if isinstance(data.get("fields"), dict) else {}

    def pick(*keys):
        """Try `data`, `header`, `fields` for the first key that exists."""
        for src in (data, header, fields):
            for k in keys:
                if k in src and src[k] is not None:
                    return src[k]
        return None

    line_items: list[dict] = []
    if isinstance(items_blk, dict):
        # Either {columns, rows} table-shape or {items: [...]} list-shape.
        cols = items_blk.get("columns") or []
        rows = items_blk.get("rows") or []
        if cols and rows:
            for r in rows:
                row = dict(zip([str(c).lower() for c in cols], r))
                line_items.append({
                    "description": row.get("description") or row.get("الوصف") or row.get("item") or "",
                    "quantity": row.get("quantity") or row.get("qty") or row.get("الكمية") or "",
                    "unit_price": row.get("unit_price") or row.get("سعر الوحدة") or row.get("price") or "",
                    "line_total": row.get("line_total") or row.get("total") or row.get("الإجمالي") or "",
                })
    elif isinstance(items_blk, list):
        for it in items_blk:
            if isinstance(it, dict):
                line_items.append({
                    "description": _val(it.get("description") or it.get("desc") or it.get("name")),
                    "quantity": _val(it.get("quantity") or it.get("qty")),
                    "unit_price": _val(it.get("unit_price") or it.get("price")),
                    "line_total": _val(it.get("line_total") or it.get("total")),
                })

    return {
        "supplier_name": _val(supplier.get("name") if supplier else pick("supplier_name", "from", "seller")),
        "supplier_tax_id": _val(supplier.get("tax_id") if supplier else pick("supplier_tax_id", "tax_id", "tin", "vat")),
        "supplier_tax_id_label": "VAT / TIN",
        "supplier_address": _val(supplier.get("address") if supplier else pick("supplier_address")),
        "supplier_phone": _val(supplier.get("phone") if supplier else pick("supplier_phone", "phone")),
        "supplier_email": _val(supplier.get("email") if supplier else pick("supplier_email", "email")),
        "customer_name": _val(customer.get("name") if customer else pick("customer_name", "to", "buyer", "bill_to")),
        "customer_tax_id": _val(customer.get("tax_id") if customer else pick("customer_tax_id")),
        "customer_address": _val(customer.get("address") if customer else pick("customer_address")),
        "invoice_number": _val(pick("invoice_number", "number", "inv_no", "ref")),
        "issue_date": _val(pick("issue_date", "invoice_date", "date")),
        "due_date": _val(pick("due_date", "payment_due")),
        "line_items": line_items,
        "subtotal": _amt(totals.get("subtotal") if totals else pick("subtotal")),
        "tax": _amt(totals.get("tax") if totals else pick("tax", "tax_total", "vat")),
        "discount": _amt(totals.get("discount") if totals else pick("discount")),
        "grand_total": _amt(totals.get("grand_total") if totals else pick("grand_total", "total", "total_amount")),
        "payment_terms": _val(pick("payment_terms", "terms", "notes")),
    }


def _flatten_prescription(data: dict) -> dict:
    """Flatten the rich prescription shape into the template's variable namespace."""
    if not isinstance(data, dict):
        return {}
    patient = data.get("patient") if isinstance(data.get("patient"), dict) else {}
    return {
        "doctor_name": _val(patient.get("doctor") or data.get("doctor_name") or data.get("doctor")),
        "specialty": _val(patient.get("specialty") or data.get("specialty")),
        "doctor_phone": _val(data.get("doctor_phone") or (data.get("phone_numbers") or [{}])[0] if data.get("phone_numbers") else ""),
        "patient_name": _val(patient.get("patient_name") or patient.get("name") or data.get("patient_name")),
        "patient_age": _val(patient.get("age") or data.get("age")),
        "patient_weight": _val(patient.get("weight") or data.get("weight_kg")),
        "issue_date": _val(patient.get("date") or data.get("issue_date") or data.get("date")),
        "diagnosis": _val(patient.get("diagnosis") or data.get("diagnosis")),
        "medications": [
            {
                "name": _val(m.get("name") or m.get("name_brand")),
                "drug_class": _val(m.get("drug_class") or m.get("name_generic")),
                "form": _val(m.get("form")),
                "dosage": _val(m.get("dosage") or m.get("dose")),
                "duration": _val(m.get("duration")),
            }
            for m in (data.get("medications") or []) if isinstance(m, dict)
        ],
        "as_needed": [
            {"name": _val(m.get("name")), "indication": _val(m.get("indication"))}
            for m in (data.get("as_needed") or data.get("as_needed_medications") or []) if isinstance(m, dict)
        ],
        "validations": [
            _val(v.get("check") if isinstance(v, dict) else v)
            for v in (data.get("validations") or []) if v
        ],
    }


def _flatten_national_id(data: dict) -> dict:
    if not isinstance(data, dict):
        return {}
    id_node = data.get("id_number") if isinstance(data.get("id_number"), dict) else None
    decoded = None
    if id_node and isinstance(id_node.get("decoded"), dict):
        decoded = id_node["decoded"]
    return {
        "id_number": _val(id_node or data.get("id_number") or data.get("national_id")),
        "issuing_country": _val(data.get("issuing_country") or data.get("country")),
        "full_name_ar": _val(data.get("full_name_ar") or data.get("name_ar") or data.get("name")),
        "full_name_en": _val(data.get("full_name_en") or data.get("name_en")),
        "date_of_birth": _val(data.get("date_of_birth") or data.get("birth_date")),
        "sex": _val(data.get("sex") or data.get("gender")),
        "address": _val(data.get("address")),
        "religion": _val(data.get("religion")),
        "marital_status": _val(data.get("marital_status")),
        "occupation": _val(data.get("occupation")),
        "issue_date": _val(data.get("issue_date")),
        "expiry_date": _val(data.get("expiry_date")),
        "decoded": decoded,
    }


def _flatten_business_card(data: dict) -> dict:
    if not isinstance(data, dict):
        return {}
    return {
        "name": _val(data.get("name") or data.get("full_name")),
        "title": _val(data.get("title") or data.get("job_title") or data.get("position")),
        "company": _val(data.get("company") or data.get("organization")),
        "phone": _val(data.get("phone") or data.get("phone_number") or data.get("mobile")),
        "email": _val(data.get("email")),
        "website": _val(data.get("website") or data.get("url")),
        "address": _val(data.get("address")),
    }


def _flatten_bank_statement(data: dict) -> dict:
    if not isinstance(data, dict):
        return {}
    account = data.get("account") if isinstance(data.get("account"), dict) else {}
    period_node = data.get("statement_period") if isinstance(data.get("statement_period"), dict) else None
    period_str = ""
    if period_node:
        period_str = f"{_val(period_node.get('from'))} → {_val(period_node.get('to'))}"
    txns = []
    for t in (data.get("transactions") or []):
        if isinstance(t, dict):
            txns.append({
                "date": _val(t.get("date") or t.get("value_date")),
                "description": _val(t.get("description") or t.get("desc")),
                "debit": _amt(t.get("debit") or {"value": t.get("debit")} if t.get("debit") is not None else None),
                "credit": _amt(t.get("credit") or {"value": t.get("credit")} if t.get("credit") is not None else None),
                "balance": _amt(t.get("balance")),
            })
    return {
        "holder_name": _val(account.get("holder_name") or data.get("holder_name")),
        "iban": _val(account.get("iban") or data.get("iban")),
        "currency": _val(account.get("currency") or data.get("currency")),
        "branch": _val(account.get("branch") or data.get("branch")),
        "period": period_str,
        "opening_balance": _amt(data.get("opening_balance")),
        "closing_balance": _amt(data.get("closing_balance")),
        "total_credits": _amt(data.get("total_credits")),
        "total_debits": _amt(data.get("total_debits")),
        "transactions": txns,
        "reconciliation": data.get("reconciliation") if isinstance(data.get("reconciliation"), dict) else None,
    }


_FLATTENERS = {
    "invoice": _flatten_invoice,
    "prescription": _flatten_prescription,
    "national_id": _flatten_national_id,
    "business_card": _flatten_business_card,
    "bank_statement": _flatten_bank_statement,
}


_ALIASES = {
    "egyptian_id": "national_id",   # Egyptian ID matches the national_id field set
    # NOTE: passport is intentionally NOT aliased here. Its field names
    # (document_number, surname, nationality, mrz, ...) don't match
    # national_id's (id_number, full_name_ar, ...), so the flattener can't
    # populate the template and you get empty chrome. Falling through to
    # the generic walker produces a usable polished page until a dedicated
    # passport template + flattener lands.
}


def _is_empty_flatten(flat: dict) -> bool:
    """Return True when a specialized flattener found nothing useful — meaning
    the data shape didn't match what the template expects. The caller should
    fall through to the generic walker so the user gets a real result instead
    of empty chrome."""
    if not isinstance(flat, dict):
        return True
    skip_keys = {"supplier_tax_id_label"}   # fixed labels, not extracted values
    populated = 0
    for k, v in flat.items():
        if k in skip_keys:
            continue
        if v in (None, "", [], {}):
            continue
        if isinstance(v, list) and not any(v):
            continue
        populated += 1
    return populated < 2


# Keyword heuristics that classify a field into a logical section. The order
# matters — we check each group in turn and assign the field to the first match.
_FIELD_GROUPS = (
    ("Identity", ("name", "full_name", "given", "surname", "father", "mother", "اسم", "الاسم")),
    ("Demographics", ("birth", "born", "age", "sex", "gender", "nationality", "religion",
                       "marital", "occupation", "profession", "ميلاد", "جنسية", "النوع",
                       "الجنس", "ديانة", "وظيفة", "مهنة")),
    ("Document Information", ("document_no", "passport_no", "id_no", "id_number", "number",
                              "issue", "expiry", "valid_until", "office", "issuing",
                              "national_id", "إصدار", "انتهاء", "رقم", "مكتب", "قومي")),
    ("Address & Contact", ("address", "city", "country", "phone", "email", "website",
                            "postal", "زip", "عنوان", "مدينة", "هاتف", "بريد")),
    ("Financial", ("total", "amount", "subtotal", "tax", "discount", "balance", "currency",
                    "iban", "swift", "إجمالي", "مبلغ", "رصيد", "ضريبة")),
    ("Machine-Readable Codes", ("mrz", "barcode", "qr")),
)

# Top-level envelope keys we never render as user-facing fields.
_SKIP_TOP = {"summary", "document_type", "layout", "language", "extracted_at",
             "validations", "warnings", "confidence", "_layout_html", "raw_text",
             "doctype_label", "cover_meta", "sections"}


def _classify_group(key: str) -> str:
    """Return the section name a field belongs in, or 'Other Details'."""
    kl = key.lower()
    for group, words in _FIELD_GROUPS:
        if any(w in kl for w in words):
            return group
    return "Other Details"


# Arabic labels for the fixed section titles + cover-metadata labels the generic
# report emits. Classification stays in English (stable keys); only the rendered
# `title`/label is localized, so nothing else in the walker needs to change.
_SECTION_AR = {
    "Identity": "الهوية",
    "Demographics": "البيانات الشخصية",
    "Document Information": "بيانات المستند",
    "Address & Contact": "العنوان والتواصل",
    "Financial": "البيانات المالية",
    "Machine-Readable Codes": "رموز قابلة للقراءة آليًا",
    "Other Details": "تفاصيل أخرى",
    "Machine-Readable Zone": "المنطقة القابلة للقراءة آليًا",
    "Validation Checks": "فحوصات التحقق",
    "Result": "النتيجة",
    # cover-metadata labels
    "Document Type": "نوع المستند",
    "Document Number": "رقم المستند",
    "Issue Date": "تاريخ الإصدار",
    "Expiry / Due Date": "تاريخ الانتهاء / الاستحقاق",
    "Issuing Office / Holder": "جهة الإصدار / الحامل",
    # multi-document report labels
    "Number of Documents": "عدد المستندات",
    "Document Types": "أنواع المستندات",
    "Extracted At": "تاريخ الاستخراج",
}


def _localize_section(title: str, lang: str) -> str:
    """Translate a fixed English section/label title to Arabic when lang=='ar'.
    Unknown titles (e.g. field-derived nested-table names) pass through as-is."""
    if lang == "ar":
        return _SECTION_AR.get(title, title)
    return title


def _is_arabic(s: str) -> bool:
    return any("؀" <= ch <= "ۿ" for ch in str(s))


def _pretty_label(key: str) -> str:
    """Turn snake_case into a Human Label, preserving Arabic words as-is."""
    if _is_arabic(key):
        return key
    return key.replace("_", " ").title()


def _build_field_entry(label: str, node):
    """Cell-shape → template dict with confidence tier."""
    v = _val(node)
    if not v:
        return None
    cls = ""
    if isinstance(node, dict) and isinstance(node.get("confidence"), (int, float)):
        c = node["confidence"]
        if c < 0.6: cls = "fail"
        elif c < 0.85: cls = "low"
    return {"label": label, "value": v, "cls": cls}


def _pair_bilingual(items: list[dict]) -> list[dict]:
    """Merge English+Arabic field pairs into one row with EN value primary +
    AR value as secondary. Heuristic: an Arabic-labeled item whose neighbour
    (within 4 positions) is its likely English counterpart → fold together.

    Keeps the visual report clean — passport extractions emit 'Full Name' +
    'الإسم' as two separate rows otherwise.
    """
    out: list[dict] = []
    used = set()
    # Build name → AR equivalent map of common pairs
    AR_PAIRS = {
        "Full Name": ("الإسم", "الاسم"),
        "Place Of Birth": ("مكان الميلاد",),
        "Date Of Birth": ("تاريخ الميلاد",),
        "Nationality": ("الجنسية",),
        "Sex": ("النوع", "الجنس"),
        "Date Of Issue": ("تاريخ الإصدار",),
        "Date Of Expiry": ("تاريخ الانتهاء",),
        "Date Of Expirey": ("تاريخ الانتهاء",),  # common typo
        "Profession": ("الوظيفة", "المهنة"),
        "Address": ("العنوان",),
    }
    by_label = {it["label"]: i for i, it in enumerate(items)}
    for i, it in enumerate(items):
        if i in used:
            continue
        pair_targets = AR_PAIRS.get(it["label"], ())
        merged = dict(it)
        for ar in pair_targets:
            if ar in by_label and by_label[ar] not in used:
                idx = by_label[ar]
                merged["value_secondary"] = items[idx]["value"]
                used.add(idx)
                break
        out.append(merged)
        used.add(i)
    return out


def _detect_mrz(data: dict) -> list[str] | None:
    """Pull MRZ lines out of the extracted data if present."""
    lines = []
    for k, v in data.items():
        kl = k.lower()
        if "mrz" in kl and "line" in kl:
            val = _val(v)
            if val: lines.append(val)
        elif kl == "mrz" and isinstance(v, dict) and isinstance(v.get("lines"), list):
            for ln in v["lines"]:
                if ln: lines.append(str(ln))
    return lines or None


def _flatten_multi(docs: list, envelope: dict, fallback_doctype: str, lang: str = "en") -> dict:
    """Render a multi-document extraction (3 passports, ID front+back, contract
    pages, etc.). Each document gets its own grouped section block."""
    # Decide a sensible doctype label: "3 Passports" if all same, "3 Documents"
    # otherwise. Use the per-doc document_type from the model.
    types = [str(d.get("document_type") or "document").lower().replace("_", " ") for d in docs]
    n = len(docs)
    if len(set(types)) == 1:
        # All same → pluralize (English-only — Arabic morphology is in the template strings)
        base = types[0].title()
        plural = base + "s" if not base.endswith("s") else base
        doctype_label = f"{n} {plural}"
    else:
        doctype_label = f"{n} Documents"

    # Cover metadata: summary across all docs
    cover_meta: list[tuple[str, str]] = [
        (_localize_section("Number of Documents", lang), str(n)),
        (_localize_section("Document Types", lang), ", ".join(sorted(set(t.title() for t in types)))),
    ]
    if envelope.get("extracted_at"):
        cover_meta.append((_localize_section("Extracted At", lang), str(envelope["extracted_at"])))

    summary = envelope.get("summary") if isinstance(envelope.get("summary"), str) else None
    if not summary:
        summary = (f"استخراج {n} مستندات — راجع كل قسم أدناه للحقول والرموز وفحوصات التحقق لكل مستند."
                   if lang == "ar" else
                   f"Extraction of {n} documents — see each section below for the per-document "
                   f"fields, machine-readable codes, and validation checks.")

    sections: list[dict] = []
    for i, doc in enumerate(docs, start=1):
        # Recursively use the single-doc walker for each document, then promote
        # its sections into a single "Document N" header here.
        per_doc = _flatten_generic(doc, str(doc.get("document_type") or fallback_doctype), lang)
        # Build a header bar for this document
        sub_label = per_doc.get("doctype_label") or "Document"
        doc_hdr = (f"المستند {i} من {n} — {sub_label}" if lang == "ar"
                   else f"Document {i} of {n} — {sub_label}")
        sections.append({
            "title": doc_hdr,
            "kind": "doc_header",
            "summary": per_doc.get("summary") or "",
            "cover_meta": per_doc.get("cover_meta") or [],
        })
        # Renumber per-doc sections so they show as 1.1, 1.2, … in the parent.
        for s in per_doc.get("sections") or []:
            # Tone the bars teal so per-doc sub-sections visually differ from
            # the top-level summary banner.
            s_copy = dict(s)
            if "tone" not in s_copy:
                s_copy["tone"] = "tone-teal"
            sections.append(s_copy)

    return {
        "doctype_label": doctype_label,
        "cover_meta": cover_meta,
        "summary": summary,
        "sections": sections,
    }


def _flatten_generic(data: dict, doctype: str, lang: str = "en") -> dict:
    """Walk an arbitrary extracted-result tree and emit a `sections` list the
    generic template can render — corporate-report style.

    The walker:
      1. Pulls cover metadata (doc number, date, status) for the title block.
      2. Classifies each leaf field into a logical group (Identity, Demographics,
         Document Info, Address, Financial, Codes, Other) so the report reads
         like a real document instead of a card wall.
      3. Pairs bilingual EN+AR fields into single rows with EN primary, AR
         secondary, to halve the visual noise.
      4. Detects MRZ / barcode / code blocks and renders them in monospace.
      5. Surfaces validations from the envelope as colored callouts.
      6. Multi-document mode: if `data.documents` is a list, renders one full
         per-document section block for each (with its own grouped fields,
         MRZ, validations), instead of a single useless meta-table.
    """
    if not isinstance(data, dict):
        return {
            "doctype_label": doctype.replace("_", " ").title(),
            "sections": [{"title": _localize_section("Result", lang), "kind": "text", "body": str(data)}],
        }

    # ── Multi-document case (3 passports, contract+addendum, ID front+back, …)
    docs = data.get("documents")
    if isinstance(docs, list) and docs and all(isinstance(d, dict) for d in docs):
        return _flatten_multi(docs, data, doctype, lang)

    doctype_label = data.get("document_type") or doctype
    doctype_label = str(doctype_label).replace("_", " ").replace("-", " ").title()
    summary = data.get("summary") if isinstance(data.get("summary"), str) else None
    mrz_lines = _detect_mrz(data)

    # ---- Cover metadata: pluck the 4-6 most useful identifiers ----
    cover_meta: list[tuple[str, str]] = []
    cover_map = [
        ("Document Type", ("document_type",)),
        ("Document Number", ("passport_no", "document_no", "id_number", "id_no",
                              "invoice_number", "number")),
        ("Issue Date", ("issue_date", "date_of_issue", "date", "invoice_date")),
        ("Expiry / Due Date", ("expiry_date", "date_of_expiry", "valid_until", "due_date")),
        ("Issuing Office / Holder", ("issuing_office", "issuing_authority",
                                      "supplier_name", "holder_name", "full_name")),
    ]
    for label, keys in cover_map:
        for k in keys:
            if k in data and data[k]:
                v = _val(data[k])
                if v:
                    if label == "Document Type":
                        v = str(v).replace("_", " ").title()
                    cover_meta.append((_localize_section(label, lang), v))
                    break

    # ---- Walk all top-level fields into groups ----
    groups: dict[str, list[dict]] = {}
    nested_tables: list[dict] = []
    nested_dicts: list[tuple[str, dict]] = []
    SKIP_KEYS_FOR_GROUPS = _SKIP_TOP | {kk for _, ks in cover_map for kk in ks}
    SKIP_KEYS_FOR_GROUPS = SKIP_KEYS_FOR_GROUPS | {"mrz_line1", "mrz_line2", "mrz",
                                                    "MRZ Line1", "MRZ Line2"}
    for k, v in data.items():
        if k in _SKIP_TOP:
            continue
        # MRZ lines already extracted; skip.
        if "mrz" in k.lower():
            continue
        if isinstance(v, dict):
            if "value" in v:
                entry = _build_field_entry(_pretty_label(k), v)
                if entry: groups.setdefault(_classify_group(k), []).append(entry)
            elif "columns" in v and "rows" in v:
                nested_tables.append({
                    "title": _pretty_label(k),
                    "kind": "table",
                    "columns": [str(c) for c in (v.get("columns") or [])],
                    "rows": [[("" if c is None else str(c)) for c in r] for r in v["rows"]],
                })
            else:
                nested_dicts.append((k, v))
        elif isinstance(v, list):
            if v and all(isinstance(x, dict) for x in v):
                cols: list[str] = []
                for item in v:
                    for kk in item.keys():
                        if kk != "confidence" and kk not in cols:
                            cols.append(kk)
                rows = [
                    [_val(item.get(c)) if isinstance(item.get(c), dict) else
                     ("" if item.get(c) is None else str(item.get(c)))
                     for c in cols]
                    for item in v
                ]
                nested_tables.append({
                    "title": _pretty_label(k),
                    "kind": "table",
                    "columns": [_pretty_label(c) for c in cols],
                    "rows": rows,
                })
            elif v and all(isinstance(x, str) for x in v):
                nested_tables.append({"title": _pretty_label(k), "kind": "list", "items": v})
        elif isinstance(v, (str, int, float)) and v not in (None, ""):
            entry = _build_field_entry(_pretty_label(k), v)
            if entry: groups.setdefault(_classify_group(k), []).append(entry)

    # ---- Emit sections in a predictable, document-y order ----
    GROUP_ORDER = ["Identity", "Demographics", "Document Information",
                    "Address & Contact", "Financial", "Other Details"]
    sections: list[dict] = []
    for gname in GROUP_ORDER:
        items = groups.pop(gname, None)
        if items:
            items = _pair_bilingual(items)
            sections.append({"title": _localize_section(gname, lang), "kind": "fields", "items": items})
    # Any new groups we didn't anticipate
    for gname, items in groups.items():
        items = _pair_bilingual(items)
        sections.append({"title": _localize_section(gname, lang), "kind": "fields", "items": items})

    # Tables and lists detected during the walk
    sections.extend(nested_tables)

    # Nested dicts (e.g. patient: {...} on Rx) become sub-sections
    for k, v in nested_dicts:
        sub_items = []
        for kk, vv in v.items():
            if kk == "confidence":
                continue
            entry = _build_field_entry(_pretty_label(kk), vv)
            if entry: sub_items.append(entry)
        if sub_items:
            sections.append({"title": _pretty_label(k), "kind": "fields",
                              "items": _pair_bilingual(sub_items)})

    # MRZ block in monospace
    if mrz_lines:
        sections.append({"title": _localize_section("Machine-Readable Zone", lang),
                          "kind": "mono", "lines": mrz_lines, "tone": "tone-teal"})

    # Validations from the envelope → colored callouts
    if isinstance(data.get("validations"), list) and data["validations"]:
        callouts = []
        for v in data["validations"]:
            if isinstance(v, dict):
                tone = "pass" if v.get("ok") else "fail"
                mark = "✓" if v.get("ok") else "✗"
                callouts.append({
                    "tone": tone,
                    "body": f"{mark}  {v.get('check', '')} — {v.get('detail', '')}",
                })
        if callouts:
            sections.append({"title": _localize_section("Validation Checks", lang),
                              "kind": "callouts", "items": callouts,
                              "tone": "tone-amber"})

    return {
        "doctype_label": doctype_label,
        "cover_meta": cover_meta,
        "summary": summary,
        "sections": sections,
    }


def render(document_type: str, data: dict, *, lang: str = "ar",
           extracted_at: str = "") -> tuple[str, str]:
    """Render the polished HTML.

    Resolution order:
      1. If document_type matches a specialized template → use it.
      2. If document_type has an alias to a specialized template → use the alias.
      3. Otherwise → use the generic walker template (works on any shape).

    Returns (html_text, doctype_used).
    """
    raw_dt = document_type or "unknown"
    dt = raw_dt.lower().replace("-", "_")
    dt = _ALIASES.get(dt, dt)
    direction = "rtl" if lang == "ar" else "ltr"

    # Path B — Gemini-generated visual reproduction takes priority over everything.
    if isinstance(data, dict) and isinstance(data.get("_layout_html"), str) \
       and data["_layout_html"].strip().lower().startswith(("<!doctype", "<html")):
        return data["_layout_html"], "visual"

    if dt in SUPPORTED:
        flatten = _FLATTENERS.get(dt)
        flat = flatten(data) if flatten else {}
        # Safety net: specialized flattener returned almost nothing → fall
        # through to the generic walker so the user gets a populated page
        # instead of empty branded chrome.
        if not _is_empty_flatten(flat):
            tmpl = _jinja.get_template(f"{dt}.html")
            html = tmpl.render(data=data, lang=lang, dir=direction,
                               extracted_at=extracted_at, **flat)
            return html, dt
    # Fallback — every doctype gets a polished output via the generic walker.
    flat = _flatten_generic(data, raw_dt, lang)
    tmpl = _jinja.get_template("generic.html")
    html = tmpl.render(data=data, lang=lang, dir=direction,
                       extracted_at=extracted_at, **flat)
    return html, "generic"


# ---- Output formats --------------------------------------------------------

def render_html(document_type: str, data: dict, lang: str = "ar",
                extracted_at: str = "") -> tuple[bytes, str, str]:
    """Polished HTML — opens in browser, prints to PDF cleanly via Ctrl+P."""
    html, dt = render(document_type, data, lang=lang, extracted_at=extracted_at)
    return html.encode("utf-8"), "text/html; charset=utf-8", f"{dt}.html"


def render_pdf(document_type: str, data: dict, lang: str = "ar",
               extracted_at: str = "") -> tuple[bytes, str, str]:
    """PDF via weasyprint — renders the SAME polished HTML to PDF, preserving
    every visual (cover hero, navy section bars, table chrome, MRZ mono block,
    callouts). Falls back to the legacy reportlab text renderer if weasyprint
    is missing on the server.
    """
    raw_dt = document_type or "document"
    name = raw_dt.lower().replace("-", "_")
    try:
        from weasyprint import HTML  # heavy import — defer
        html, dt_used = render(document_type, data, lang=lang,
                                 extracted_at=extracted_at)
        pdf_bytes = HTML(string=html).write_pdf()
        return (pdf_bytes, "application/pdf",
                f"{dt_used if dt_used != 'visual' else name}_reproduced.pdf")
    except ImportError:
        # Legacy fallback: structured plaintext via reportlab. Ugly but works.
        return _render_pdf_text_fallback(document_type, data, lang, extracted_at)


def _render_pdf_text_fallback(document_type: str, data: dict, lang: str,
                                extracted_at: str) -> tuple[bytes, str, str]:
    """Last-resort PDF renderer when weasyprint isn't installed."""
    from app import exports as _exports
    raw_dt = document_type or "document"
    dt = raw_dt.lower().replace("-", "_")
    dt = _ALIASES.get(dt, dt)
    flat = _FLATTENERS[dt](data) if dt in _FLATTENERS else None
    if flat is not None and not _is_empty_flatten(flat):
        text_lines = _plaintext_from_flat(dt, flat, lang=lang)
    else:
        gen = _flatten_generic(data, raw_dt, lang)
        text_lines = [gen.get("doctype_label", "Document"), ""]
        if gen.get("summary"):
            text_lines += [gen["summary"], ""]
        for section in gen.get("sections", []):
            text_lines.append(f"--- {section.get('title','')} ---")
            if section.get("kind") == "fields":
                for f in section.get("items", []):
                    text_lines.append(f"  {f.get('label','')}: {f.get('value','')}")
            text_lines.append("")
    name = dt if dt in _FLATTENERS else "document"
    body, mt, fn = _exports.text_to_pdf("\n".join(text_lines),
                                         title=f"{name}_reproduced")
    return body, mt, f"{name}_reproduced.pdf"


def render_docx(document_type: str, data: dict, lang: str = "ar",
                extracted_at: str = "") -> tuple[bytes, str, str]:
    """Word doc — uses python-docx for proper structured output. Any doctype
    works: specialized renderer for invoice, generic table for everything else."""
    import io as _io
    from docx import Document

    raw_dt = document_type or "document"
    dt = raw_dt.lower().replace("-", "_")
    dt = _ALIASES.get(dt, dt)

    doc = Document()
    flat = _FLATTENERS[dt](data) if dt in _FLATTENERS else None
    use_generic = flat is None or _is_empty_flatten(flat)
    if not use_generic and dt == "invoice":
        _docx_invoice(doc, flat, lang)
    elif not use_generic:
        _docx_generic(doc, flat, lang)
    else:
        # No flattener OR specialized one returned empty — walk the tree.
        gen = _flatten_generic(data, raw_dt, lang)
        _docx_from_sections(doc, gen, lang)
        dt = "document"
    buf = _io.BytesIO()
    doc.save(buf)
    return (buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{dt if dt != 'generic' else 'document'}_reproduced.docx")


def _docx_from_sections(doc, gen: dict, lang: str) -> None:
    """Render the generic-walker sections list into a Word doc — corporate style."""
    from docx.shared import Pt, RGBColor

    # ── Big title (cover block) ──
    title = doc.add_heading(gen.get("doctype_label") or
                             ("تقرير مستند" if lang == "ar" else "Document Report"),
                             level=0)
    sub = doc.add_paragraph()
    sub_run = sub.add_run("DATA EXTRACTION REPORT" if lang != "ar" else "تقرير استخراج البيانات")
    sub_run.bold = True; sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0xb4, 0x53, 0x09)

    # ── Cover metadata table ──
    cover_meta = gen.get("cover_meta") or []
    if cover_meta:
        meta_tbl = doc.add_table(rows=1, cols=2)
        meta_tbl.style = "Light Grid Accent 5"
        hdr = meta_tbl.rows[0].cells
        hdr[0].text = "Field" if lang != "ar" else "الحقل"
        hdr[1].text = "Value" if lang != "ar" else "القيمة"
        for label, value in cover_meta:
            r = meta_tbl.add_row().cells
            r[0].text = label
            r[1].text = str(value)

    # ── Summary ──
    if gen.get("summary"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(gen["summary"])
        r.italic = True; r.font.size = Pt(11)

    # ── Sections ──
    for i, section in enumerate(gen.get("sections", []), start=1):
        h = doc.add_heading(f"{i}. {section.get('title','')}", level=1)
        kind = section.get("kind")
        if kind == "fields":
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light Grid Accent 1"
            for f in section.get("items", []):
                row = tbl.add_row().cells
                row[0].text = f.get("label", "")
                v = str(f.get("value", ""))
                if f.get("value_secondary"):
                    v += "\n" + str(f["value_secondary"])
                row[1].text = v
        elif kind == "table":
            cols = section.get("columns", [])
            tbl = doc.add_table(rows=1, cols=max(1, len(cols)))
            tbl.style = "Light Grid Accent 1"
            for j, c in enumerate(cols):
                tbl.rows[0].cells[j].text = str(c)
            for r in section.get("rows", []):
                row = tbl.add_row().cells
                for j, v in enumerate(r):
                    if j < len(row):
                        row[j].text = "" if v is None else str(v)
        elif kind == "mono":
            for line in section.get("lines", []):
                p = doc.add_paragraph()
                run = p.add_run(str(line))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
        elif kind == "callouts":
            for c in section.get("items", []):
                p = doc.add_paragraph()
                run = p.add_run(c.get("body", ""))
                if c.get("tone") == "pass":
                    run.font.color.rgb = RGBColor(0x06, 0x5f, 0x46)
                elif c.get("tone") == "fail":
                    run.font.color.rgb = RGBColor(0x99, 0x1b, 0x1b)
        elif kind == "doc_header":
            # Per-document mini-cover inside a multi-doc report.
            if section.get("summary"):
                p = doc.add_paragraph()
                r = p.add_run(section["summary"])
                r.italic = True
            cm = section.get("cover_meta") or []
            if cm:
                meta_tbl = doc.add_table(rows=0, cols=2)
                meta_tbl.style = "Light Grid Accent 5"
                for label, value in cm:
                    r = meta_tbl.add_row().cells
                    r[0].text = str(label); r[1].text = str(value)
        elif kind == "list":
            for item in section.get("items", []):
                doc.add_paragraph(str(item), style="List Bullet")
        elif kind == "text":
            doc.add_paragraph(section.get("body", ""))


# ---- Helpers ---------------------------------------------------------------

def _plaintext_from_flat(doctype: str, flat: dict, lang: str) -> list[str]:
    """Turn the flattened dict into structured plaintext lines for PDF."""
    if doctype == "invoice":
        out = [
            f"INVOICE #{flat.get('invoice_number','')}",
            f"Date: {flat.get('issue_date','')}",
            "",
            "FROM",
            flat.get("supplier_name", ""),
            f"Tax ID: {flat.get('supplier_tax_id','')}",
            flat.get("supplier_address", ""),
            "",
            "BILL TO",
            flat.get("customer_name", ""),
            flat.get("customer_address", ""),
            "",
            "LINE ITEMS",
        ]
        for it in flat.get("line_items", []):
            out.append(f"  {it.get('description','')} | qty={it.get('quantity','')} | unit={it.get('unit_price','')} | total={it.get('line_total','')}")
        out += [
            "",
            f"Subtotal:    {flat.get('subtotal','')}",
            f"Tax:         {flat.get('tax','')}",
            f"GRAND TOTAL: {flat.get('grand_total','')}",
        ]
        return out
    return [f"{k}: {v}" for k, v in flat.items() if v]


def _docx_invoice(doc, flat: dict, lang: str) -> None:
    """Render the invoice into a Word doc with header / parties / items / totals."""
    from docx.shared import Pt
    title = doc.add_heading(
        ("فاتورة" if lang == "ar" else "INVOICE") + f"  #{flat.get('invoice_number','')}",
        level=0,
    )
    p = doc.add_paragraph()
    p.add_run(("التاريخ: " if lang == "ar" else "Date: ") + str(flat.get("issue_date", "")))
    if flat.get("due_date"):
        p.add_run(("    الاستحقاق: " if lang == "ar" else "    Due: ") + str(flat["due_date"]))

    # From / To
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light List"
    c0, c1 = tbl.rows[0].cells
    c0.text = ("من" if lang == "ar" else "FROM") + "\n" + str(flat.get("supplier_name", ""))
    if flat.get("supplier_tax_id"):
        c0.text += "\n" + ("الرقم الضريبي: " if lang == "ar" else "Tax ID: ") + str(flat["supplier_tax_id"])
    if flat.get("supplier_address"):
        c0.text += "\n" + str(flat["supplier_address"])
    c1.text = ("إلى" if lang == "ar" else "BILL TO") + "\n" + str(flat.get("customer_name", ""))
    if flat.get("customer_tax_id"):
        c1.text += "\n" + ("الرقم الضريبي: " if lang == "ar" else "Tax ID: ") + str(flat["customer_tax_id"])

    # Line items
    items = flat.get("line_items") or []
    if items:
        doc.add_paragraph()
        items_tbl = doc.add_table(rows=1, cols=4)
        items_tbl.style = "Light Grid Accent 1"
        hdr = items_tbl.rows[0].cells
        if lang == "ar":
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "الوصف", "الكمية", "سعر الوحدة", "الإجمالي"
        else:
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Description", "Qty", "Unit price", "Total"
        for it in items:
            row = items_tbl.add_row().cells
            row[0].text = str(it.get("description", ""))
            row[1].text = str(it.get("quantity", ""))
            row[2].text = str(it.get("unit_price", ""))
            row[3].text = str(it.get("line_total", ""))

    # Totals
    doc.add_paragraph()
    totals_tbl = doc.add_table(rows=0, cols=2)
    for label, key, bold in (
        (("الإجمالي الفرعي" if lang == "ar" else "Subtotal"), "subtotal", False),
        (("الضريبة" if lang == "ar" else "Tax"), "tax", False),
        (("الخصم" if lang == "ar" else "Discount"), "discount", False),
        (("الإجمالي الكلي" if lang == "ar" else "GRAND TOTAL"), "grand_total", True),
    ):
        v = flat.get(key)
        if not v:
            continue
        row = totals_tbl.add_row().cells
        row[0].text = label
        row[1].text = str(v)
        if bold:
            for r in row:
                for para in r.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(13)


def _docx_generic(doc, flat: dict, lang: str) -> None:
    """Fallback: dump the flattened fields into a 2-column table."""
    doc.add_heading(("نسخة منسّقة" if lang == "ar" else "Reproduced Document"), level=0)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    for k, v in flat.items():
        if v in (None, "", []):
            continue
        row = tbl.add_row().cells
        row[0].text = str(k).replace("_", " ").title()
        row[1].text = str(v) if not isinstance(v, list) else "; ".join(str(x) for x in v)
