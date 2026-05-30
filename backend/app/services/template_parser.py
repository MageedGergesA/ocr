"""Template introspection — read an uploaded template file (Excel, Word, PDF
form, HTML) and extract its field structure as `{field_name: human_description}`.

The output is the same schema shape Mostakhles already uses for saved templates
and `extract_schema`, so a parsed template plugs straight into the existing
extraction pipeline.

Strategy per format:
- xlsx   : openpyxl reads row 1 of the first sheet as column headers
- docx   : python-docx scans the first table's header row, content controls,
           and "Label:" patterns in paragraphs
- pdf    : pdfplumber pulls AcroForm field names; falls back to Gemini-based
           layout introspection if the PDF has no form fields
- html   : BeautifulSoup walks <input>, <select>, <textarea> + their <label>s

Each parser returns: list[tuple[name, description]]. Callers build the schema
dict from that.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Iterable

log = logging.getLogger("mostakhles")

# Map common content_type / extension hints to our internal kind.
CONTENT_TYPE_TO_KIND = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",  # legacy xls — openpyxl may struggle but try
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "application/pdf": "pdf",
    "text/html": "html",
    "application/xhtml+xml": "html",
}


def detect_kind(filename: str | None, content_type: str | None) -> str | None:
    if content_type and content_type in CONTENT_TYPE_TO_KIND:
        return CONTENT_TYPE_TO_KIND[content_type]
    if not filename:
        return None
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    return {
        "xlsx": "xlsx", "xls": "xlsx",
        "docx": "docx", "doc": "docx",
        "pdf": "pdf",
        "html": "html", "htm": "html",
    }.get(ext)


def _clean_header(s: str) -> str:
    """Turn a column header like 'Vendor Name (مطلوب) *' into a clean field key
    'vendor_name'. Keeps Arabic letters; strips parens/asterisks/whitespace;
    lowercases English so introspect ↔ fill matching is case-insensitive."""
    s = (s or "").strip()
    # Strip trailing decorators
    s = re.sub(r"\s*\([^)]*\)\s*$", "", s)
    s = s.rstrip("*: \t")
    # Replace whitespace and most punctuation with underscores; keep Arabic + word chars
    s = re.sub(r"[\s/\\\-]+", "_", s)
    s = re.sub(r"[^\w؀-ۿ_]", "", s)
    return s.lower() or "field"


def parse_xlsx(file_bytes: bytes) -> list[tuple[str, str]]:
    """Read row 1 of the first non-empty worksheet as column headers."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return []
    try:
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        for ws in wb.worksheets:
            row1 = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
            if not row1:
                continue
            headers = [str(c).strip() for c in row1 if c is not None and str(c).strip()]
            if headers:
                return [(_clean_header(h), h) for h in headers]
        return []
    except Exception as e:  # noqa: BLE001
        log.warning("parse_xlsx failed: %s", e)
        return []


def parse_docx(file_bytes: bytes) -> list[tuple[str, str]]:
    """Discover fields in a Word template:
      1) Header row of the first table (table-based forms — most common pattern)
      2) Content controls (sdt elements) — Word's modern form fields
      3) Paragraphs matching 'Label:' or 'Label :' patterns
    """
    try:
        from docx import Document
    except ImportError:
        return []
    found: list[tuple[str, str]] = []
    seen = set()

    def add(name: str, desc: str | None = None):
        key = _clean_header(name)
        if key and key not in seen:
            seen.add(key)
            found.append((key, (desc or name).strip()))

    try:
        doc = Document(io.BytesIO(file_bytes))
        # (1) First table header row
        if doc.tables:
            hdr = doc.tables[0].rows[0]
            for cell in hdr.cells:
                txt = cell.text.strip()
                if txt:
                    add(txt)
        # (2) Content controls — Word's form fields live in <w:sdt> elements;
        # python-docx exposes them via the XML if not as objects.
        try:
            from docx.oxml.ns import qn
            for sdt in doc.element.body.iter(qn("w:sdt")):
                alias = sdt.find(qn("w:sdtPr") + "/" + qn("w:alias"))
                tag = sdt.find(qn("w:sdtPr") + "/" + qn("w:tag"))
                name = (alias.get(qn("w:val")) if alias is not None else None) \
                    or (tag.get(qn("w:val")) if tag is not None else None)
                if name:
                    add(name)
        except Exception:  # noqa: BLE001
            pass
        # (3) Paragraphs with "Label:" pattern (catches simple Arabic+English templates)
        label_re = re.compile(r"^(.{2,60}?)\s*[:：]\s*$")
        for p in doc.paragraphs:
            m = label_re.match(p.text.strip())
            if m:
                add(m.group(1))
        return found
    except Exception as e:  # noqa: BLE001
        log.warning("parse_docx failed: %s", e)
        return found


def parse_pdf(file_bytes: bytes) -> list[tuple[str, str]]:
    """Extract AcroForm field names. If no form fields are present, falls back
    to scraping likely field labels from the first page text."""
    found: list[tuple[str, str]] = []
    seen = set()

    def add(name: str, desc: str | None = None):
        key = _clean_header(name)
        if key and key not in seen:
            seen.add(key)
            found.append((key, (desc or name).strip()))

    try:
        # AcroForm fields via pypdf (already a dep)
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        fields = reader.get_fields() or {}
        for name, info in fields.items():
            if name:
                add(str(name), info.get("/TU") if isinstance(info, dict) else None)
        if found:
            return found
    except Exception as e:  # noqa: BLE001
        log.warning("parse_pdf AcroForm step failed: %s", e)

    try:
        # Fallback: scrape labels from first page text via pdfplumber
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                return found
            text = pdf.pages[0].extract_text() or ""
            label_re = re.compile(r"^([\w؀-ۿ \-/]{2,40}?)\s*[:：]")
            for line in text.splitlines():
                m = label_re.match(line.strip())
                if m:
                    add(m.group(1))
        return found
    except Exception as e:  # noqa: BLE001
        log.warning("parse_pdf fallback failed: %s", e)
        return found


def parse_html(file_bytes: bytes) -> list[tuple[str, str]]:
    """Walk form controls; prefer the <label for=...> text, fall back to name attr."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return []
    found: list[tuple[str, str]] = []
    seen = set()

    def add(name: str, desc: str | None = None):
        key = _clean_header(name)
        if key and key not in seen:
            seen.add(key)
            found.append((key, (desc or name).strip()))

    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        soup = BeautifulSoup(text, "html.parser")
        # Build id → label-text map
        labels = {}
        for lbl in soup.find_all("label"):
            target = lbl.get("for")
            if target:
                labels[target] = lbl.get_text(strip=True)
        for tag in soup.find_all(["input", "select", "textarea"]):
            if (tag.get("type") or "").lower() in {"hidden", "submit", "button", "reset"}:
                continue
            name = tag.get("name") or tag.get("id") or ""
            desc = labels.get(tag.get("id") or "") or tag.get("placeholder") or name
            if name:
                add(name, desc)
        # Also pick up table header cells (<th>) for invoice-style templates
        for table in soup.find_all("table"):
            ths = table.find_all("th")
            for th in ths:
                txt = th.get_text(strip=True)
                if txt:
                    add(txt)
            if ths:
                break  # first table only
        return found
    except Exception as e:  # noqa: BLE001
        log.warning("parse_html failed: %s", e)
        return found


def introspect(file_bytes: bytes, filename: str | None,
               content_type: str | None) -> dict:
    """Top-level entry. Returns:
        {"kind": "xlsx|docx|pdf|html", "fields": {name: desc, ...},
         "field_count": int, "warning": optional str}
    """
    kind = detect_kind(filename, content_type)
    if not kind:
        return {"kind": None, "fields": {}, "field_count": 0,
                "warning": "unsupported template type"}
    parser = {"xlsx": parse_xlsx, "docx": parse_docx,
              "pdf": parse_pdf, "html": parse_html}[kind]
    pairs = parser(file_bytes)
    fields = {k: v for k, v in pairs}
    return {"kind": kind, "fields": fields, "field_count": len(fields),
            "warning": None if fields else "no fields detected — file may be empty or use a custom layout"}
