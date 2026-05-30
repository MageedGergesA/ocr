"""Template filling — take the original template + extracted values and produce
a downloadable filled copy.

Two-format MVP: xlsx and docx (the dominant business-template formats in MENA).
PDF AcroForm fill is doable with pypdf 4.x but quirky; deferred. HTML fill is
trivial (template engine) but rarely useful — also deferred.

Strategy:
- xlsx : openpyxl reads the original template, walks the first sheet's row-1
         headers, writes a single result row at row 2 mapping by `_clean_header`.
         (For invoice-style templates with line items, callers should pass a
         list of dicts via `rows` and we add one row each.)
- docx : python-docx fills two surfaces:
           1) Cells in the second row of the first table whose row-1 header
              matches a field name (single-record templates)
           2) Content controls (<w:sdt>) — write the value to <w:t>
"""
from __future__ import annotations

import io
import logging
from typing import Any

from .template_parser import _clean_header

log = logging.getLogger("mostakhles")


def _value_of(v: Any) -> str:
    """Coerce an extracted-value dict {value, confidence} or primitive into a str."""
    if v is None:
        return ""
    if isinstance(v, dict) and "value" in v:
        return _value_of(v["value"])
    if isinstance(v, (list, tuple)):
        return ", ".join(_value_of(x) for x in v)
    return str(v)


def _flatten_to_pairs(data: dict) -> dict:
    """Accept any shape Mostakhles emits and flatten to {clean_key: str_value}.
    Supports: {field: {value, confidence}}, {patient: {field: {...}}, medications: [...]},
    {header, totals}, plain dicts.

    Wrapper keys like 'fields', 'data', 'patient', 'header', 'totals' are TRANSPARENT —
    we descend into them without prefixing. Otherwise an Excel header 'Customer Name'
    (cleaned to 'customer_name') would never match the flattened 'fields_customer_name'
    and the fill would silently leave the cell empty.
    """
    out: dict[str, str] = {}
    # Structural wrappers — descend through them without contributing to the key path.
    transparent = {"fields", "data", "patient", "header", "totals", "result", "info"}
    skip = {"confidence", "layout", "document_type", "output_lang", "files_count"}

    def walk(prefix: str, obj: Any):
        if obj is None:
            return
        if isinstance(obj, dict):
            if "value" in obj and not isinstance(obj.get("value"), (dict, list)):
                out[_clean_header(prefix)] = _value_of(obj.get("value"))
                return
            for k, v in obj.items():
                if k in skip:
                    continue
                if k in transparent:
                    walk(prefix, v)  # descend without adding to the key path
                    continue
                child = k if not prefix else f"{prefix}_{k}"
                walk(child, v)
        elif isinstance(obj, list):
            # Lists become joined strings unless they're complex (line items).
            if all(isinstance(x, (str, int, float)) for x in obj):
                out[_clean_header(prefix)] = _value_of(obj)
            # Complex lists left to caller via `rows` arg.
        else:
            out[_clean_header(prefix)] = _value_of(obj)

    walk("", data or {})
    # Also publish every key with its lowercased form so case-mismatch is forgiving.
    out.update({k.lower(): v for k, v in list(out.items())})
    return out


def fill_xlsx(template_bytes: bytes, extracted: dict,
              rows: list[dict] | None = None) -> bytes:
    """Open the template xlsx, write one data row under matching headers (single
    record). If `rows` is provided, write multiple rows (line-items / table case)."""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(template_bytes))
    ws = wb.active
    if ws is None or ws.max_row < 1:
        out = io.BytesIO(); wb.save(out); return out.getvalue()
    # Build header index: column letter index -> field key
    headers = []
    for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1), []), start=1):
        if cell.value is None:
            headers.append(None)
            continue
        headers.append(_clean_header(str(cell.value)))

    flat = _flatten_to_pairs(extracted)

    def write_row(row_num: int, src: dict):
        for col_idx, key in enumerate(headers, start=1):
            if not key:
                continue
            v = src.get(key) or src.get(key.lower()) or ""
            ws.cell(row=row_num, column=col_idx, value=v or "")

    if rows:
        for i, r in enumerate(rows, start=2):
            write_row(i, _flatten_to_pairs(r) if isinstance(r, dict) else {})
    else:
        write_row(2, flat)

    out = io.BytesIO(); wb.save(out); return out.getvalue()


def fill_docx(template_bytes: bytes, extracted: dict,
              rows: list[dict] | None = None) -> bytes:
    """Fill a Word template:
       - For the first table, write one row per record. With `rows=None` we
         write a single data row (row 2). With `rows=[{...}, {...}]` we write
         one row per record (auto-adds rows as needed).
       - For each content control (sdt), if its tag/alias matches a field, set
         its text run (single-record only).
       - For "{{ field }}" Jinja-style placeholders in paragraphs, simple replace.
    """
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(io.BytesIO(template_bytes))
    flat = _flatten_to_pairs(extracted)

    # (1) First table — write one row per record under the row-1 headers.
    if doc.tables:
        tbl = doc.tables[0]
        if len(tbl.rows) >= 1:
            header_cells = [c.text.strip() for c in tbl.rows[0].cells]
            header_keys = [_clean_header(h) for h in header_cells]
            # Multi-record mode: write one row per `rows` entry.
            records = rows if rows else [extracted]
            while len(tbl.rows) < len(records) + 1:
                tbl.add_row()
            for r_idx, record in enumerate(records, start=1):
                if r_idx >= len(tbl.rows):
                    break
                rflat = _flatten_to_pairs(record) if record is not extracted else flat
                data_row = tbl.rows[r_idx]
                for col_idx, key in enumerate(header_keys):
                    if col_idx >= len(data_row.cells) or not key:
                        continue
                    data_row.cells[col_idx].text = rflat.get(key, "")

    # (2) Content controls (sdt) → fill the inner <w:t>
    try:
        for sdt in doc.element.body.iter(qn("w:sdt")):
            alias = sdt.find(qn("w:sdtPr") + "/" + qn("w:alias"))
            tag = sdt.find(qn("w:sdtPr") + "/" + qn("w:tag"))
            name = (alias.get(qn("w:val")) if alias is not None else None) \
                or (tag.get(qn("w:val")) if tag is not None else None)
            if not name:
                continue
            key = _clean_header(name)
            value = flat.get(key)
            if value is None:
                continue
            # Find <w:t> inside the sdt content and replace its text
            content = sdt.find(qn("w:sdtContent"))
            if content is None:
                continue
            t = content.find(".//" + qn("w:t"))
            if t is not None:
                t.text = value
    except Exception as e:  # noqa: BLE001
        log.warning("fill_docx sdt step failed: %s", e)

    # (3) Simple {{ field }} placeholder substitution in paragraphs
    import re
    placeholder_re = re.compile(r"\{\{\s*([\w؀-ۿ_]+)\s*\}\}")
    for p in doc.paragraphs:
        text = p.text
        if "{{" not in text:
            continue

        def sub(m):
            return flat.get(_clean_header(m.group(1)), m.group(0))
        new = placeholder_re.sub(sub, text)
        if new != text:
            # Replace the whole paragraph's text — rough but reliable enough for
            # simple templates. Power users should use content controls.
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = new
            else:
                p.add_run(new)

    out = io.BytesIO(); doc.save(out); return out.getvalue()


def fill_template(template_bytes: bytes, kind: str, extracted: dict,
                  rows: list[dict] | None = None) -> tuple[bytes, str, str]:
    """Dispatch to the right filler. Returns (bytes, media_type, suggested_filename)."""
    if kind == "xlsx":
        return (fill_xlsx(template_bytes, extracted, rows=rows),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "extracted.xlsx")
    if kind == "docx":
        return (fill_docx(template_bytes, extracted, rows=rows),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "extracted.docx")
    raise ValueError(f"fill_template: unsupported kind '{kind}' (xlsx/docx only)")
