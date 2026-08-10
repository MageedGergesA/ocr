"""Export extracted rows to CSV / Excel / Word / PDF (Arabic-safe).

Each exporter takes `rows` = [{"page", "field", "value", "confidence"}, ...]
and returns (bytes, media_type, filename).
"""
import csv
import io

ARABIC_FONT = "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf"


def _conf_str(conf) -> str:
    return f"{round(conf * 100)}%" if isinstance(conf, (int, float)) else ""


# The four fixed column headers, per export language. `lang="en"` gives an
# English, left-to-right sheet; anything else keeps the Arabic RTL defaults.
def _labels(lang):
    if lang == "en":
        return ("Page", "Field", "Value", "Confidence")
    return ("صفحة", "الحقل", "القيمة", "الثقة")


def _report_title(lang):
    return "Extraction Result" if lang == "en" else "نتيجة الاستخراج"


def to_csv(rows, title="extraction", lang="ar"):
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(list(_labels(lang)))
    for r in rows:
        w.writerow([r.get("page"), r.get("field"), r.get("value"), _conf_str(r.get("confidence"))])
    data = ("﻿" + buf.getvalue()).encode("utf-8")  # BOM so Excel reads Arabic
    return data, "text/csv; charset=utf-8", f"{title}.csv"


def to_xlsx(rows, title="extraction", lang="ar"):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Extraction"
    ws.sheet_view.rightToLeft = (lang != "en")
    ws.append(list(_labels(lang)))
    for r in rows:
        val = "" if r.get("value") is None else str(r.get("value"))
        ws.append([r.get("page"), r.get("field"), val, _conf_str(r.get("confidence"))])
    buf = io.BytesIO()
    wb.save(buf)
    return (buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            f"{title}.xlsx")


def to_docx(rows, title="extraction", lang="ar"):
    from docx import Document

    doc = Document()
    doc.add_heading(_report_title(lang), level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    labels = _labels(lang)
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = labels
    for r in rows:
        c = table.add_row().cells
        c[0].text = str(r.get("page", ""))
        c[1].text = str(r.get("field", ""))
        c[2].text = "" if r.get("value") is None else str(r.get("value"))
        c[3].text = _conf_str(r.get("confidence"))
    buf = io.BytesIO()
    doc.save(buf)
    return (buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{title}.docx")


_FONT_READY = False


def _ar(text):
    """Shape + reorder Arabic + escape XML-special chars so the reportlab
    Paragraph parser (which interprets <b>/<i> tags) doesn't crash on raw
    `<`, `&`, or unbalanced tags coming from extracted document text
    (very common: MRZ codes contain `<<<<`, tax IDs contain `<`, etc.)."""
    import arabic_reshaper
    from bidi.algorithm import get_display
    raw = "" if text is None else str(text)
    shaped = get_display(arabic_reshaper.reshape(raw))
    # Escape AFTER shaping so the BiDi reorder works on the original glyphs.
    return (shaped
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))


def to_pdf(rows, title="extraction", lang="ar"):
    global _FONT_READY
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    if not _FONT_READY:
        pdfmetrics.registerFont(TTFont("Naskh", ARABIC_FONT))
        _FONT_READY = True

    en = lang == "en"
    align = 0 if en else 2  # 0=left (LTR), 2=right (RTL)
    labels = _labels(lang)
    head = ParagraphStyle("h", fontName="Naskh", fontSize=15, alignment=align)
    cell = ParagraphStyle("c", fontName="Naskh", fontSize=10, alignment=align,
                          wordWrap=None if en else "RTL", leading=14)

    def P(text):
        # English: plain text (no BiDi reshaping/reorder). Arabic: shape + reorder.
        raw = "" if text is None else str(text)
        if en:
            raw = raw.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            return Paragraph(raw, cell)
        return Paragraph(_ar(text), cell)

    if en:
        # LTR order: Page | Field | Value | Confidence
        table_data = [[P(labels[0]), P(labels[1]), P(labels[2]), P(labels[3])]]
        for r in rows:
            table_data.append([
                P(r.get("page")), P(r.get("field")),
                P(r.get("value")), P(_conf_str(r.get("confidence"))),
            ])
        col_widths = [45, 140, 215, 55]
    else:
        # RTL order (reversed): Confidence | Value | Field | Page
        table_data = [[P(labels[3]), P(labels[2]), P(labels[1]), P(labels[0])]]
        for r in rows:
            table_data.append([
                P(_conf_str(r.get("confidence"))), P(r.get("value")),
                P(r.get("field")), P(r.get("page")),
            ])
        col_widths = [55, 215, 140, 45]

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=title)
    table = Table(table_data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fb")]),
    ]))
    heading = "Extraction Result — Mostakhles" if en else _ar("نتيجة الاستخراج — مُستخلِص")
    doc.build([Paragraph(heading, head), Spacer(1, 12), table])
    return buf.getvalue(), "application/pdf", f"{title}.pdf"


def to_xml(rows, title="extraction", lang="ar"):
    from xml.sax.saxutils import escape
    out = ['<?xml version="1.0" encoding="UTF-8"?>', "<extraction>"]
    for r in rows:
        val = "" if r.get("value") is None else str(r.get("value"))
        out += [
            "  <field>",
            f"    <page>{r.get('page', '')}</page>",
            f"    <name>{escape(str(r.get('field', '')))}</name>",
            f"    <value>{escape(val)}</value>",
            f"    <confidence>{r.get('confidence', '')}</confidence>",
            "  </field>",
        ]
    out.append("</extraction>")
    return "\n".join(out).encode("utf-8"), "application/xml; charset=utf-8", f"{title}.xml"


EXPORTERS = {"csv": to_csv, "xlsx": to_xlsx, "docx": to_docx, "pdf": to_pdf, "xml": to_xml}


# ---- "Fields-as-columns" exports (one document per row) ----
# Better for downstream accounting / CRM imports than the legacy row-per-field
# shape: the field names become column headers, values fill a single data row.
# Multi-document results (auto_multi mode) produce one row per document.

def _flatten_for_columns(rows):
    """Turn legacy [{page, field, value, confidence}] rows into a {field: value}
    dict (one row per document) or list-of-dicts (multi-document)."""
    if not rows:
        return [{}]
    # Detect multi-document by looking for repeated field names with different pages.
    pages = {r.get("page") for r in rows if r.get("page") is not None}
    if len(pages) > 1:
        grouped: dict = {}
        for r in rows:
            grouped.setdefault(r.get("page"), {})[str(r.get("field"))] = r.get("value")
        return [{"_doc": p, **fields} for p, fields in grouped.items()]
    return [{str(r.get("field")): r.get("value") for r in rows}]


def to_csv_columns(rows, title="extraction", lang="ar"):
    """CSV with field names as column headers, values as a single data row.
    `lang` is accepted for a uniform exporter signature but unused here — the
    headers are the source field names, not fixed localizable labels."""
    docs = _flatten_for_columns(rows)
    all_fields: list[str] = []
    for d in docs:
        for k in d.keys():
            if k not in all_fields:
                all_fields.append(k)
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(all_fields)
    for d in docs:
        w.writerow([("" if d.get(k) is None else str(d.get(k))) for k in all_fields])
    data = ("﻿" + buf.getvalue()).encode("utf-8")
    return data, "text/csv; charset=utf-8", f"{title}.csv"


def to_xlsx_columns(rows, title="extraction", lang="ar"):
    """Excel with field names as column headers, values as data rows.
    `lang` accepted for a uniform signature; headers are source field names."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    docs = _flatten_for_columns(rows)
    all_fields: list[str] = []
    for d in docs:
        for k in d.keys():
            if k not in all_fields:
                all_fields.append(k)
    wb = Workbook()
    ws = wb.active
    ws.title = "Extraction"
    ws.sheet_view.rightToLeft = (lang != "en")
    # Header row
    hdr_font = Font(bold=True)
    hdr_fill = PatternFill("solid", fgColor="D1FAE5")
    for c, name in enumerate(all_fields, start=1):
        cell = ws.cell(row=1, column=c, value=str(name))
        cell.font = hdr_font
        cell.fill = hdr_fill
    # Data rows
    for r, d in enumerate(docs, start=2):
        for c, name in enumerate(all_fields, start=1):
            v = d.get(name)
            ws.cell(row=r, column=c, value=("" if v is None else str(v)))
    # Auto-size columns to ~20 chars each (good default for Arabic).
    for c in range(1, len(all_fields) + 1):
        ws.column_dimensions[chr(64 + c) if c <= 26 else "AA"].width = 22
    buf = io.BytesIO()
    wb.save(buf)
    return (buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            f"{title}.xlsx")


# Register the column-shape variants so the API can offer both.
EXPORTERS["csv_columns"] = to_csv_columns
EXPORTERS["xlsx_columns"] = to_xlsx_columns


# ---- Layout-aware Excel ----
# Take a rich extraction result (with `layout` and per-layout fields) and render
# an Excel that matches the document's shape, instead of one flat row-per-field
# table. Falls back to `to_xlsx` for unknown layouts.

def _val(node):
    """Read either a plain value or a {value, confidence} cell."""
    if isinstance(node, dict) and "value" in node:
        return node.get("value")
    return node


def _conf(node):
    if isinstance(node, dict) and "confidence" in node:
        return _conf_str(node.get("confidence"))
    return ""


def to_xlsx_layout_aware(result: dict, title: str = "extraction"):
    """Render a layout-shaped Excel based on result.layout.

    Supported layouts:
      - prescription: patient header + medications table + as-needed + validations
      - mixed (invoice/receipt): header + line items + totals
      - table: a single table sheet
      - everything else: falls back to flat rows.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    layout = (result or {}).get("layout") or ""
    doc_type = (result or {}).get("document_type") or "extraction"

    if layout == "prescription":
        wb = Workbook()
        ws = wb.active
        ws.title = "Prescription"
        ws.sheet_view.rightToLeft = True
        title_font = Font(bold=True, size=13)
        hdr_font = Font(bold=True)
        hdr_fill = PatternFill("solid", fgColor="D1FAE5")

        ws["A1"] = "روشتة / Prescription"
        ws["A1"].font = title_font

        # Patient section
        row = 3
        ws.cell(row=row, column=1, value="بيانات المريض / Patient").font = hdr_font
        row += 1
        patient = result.get("patient") or {}
        for key, node in patient.items():
            ws.cell(row=row, column=1, value=str(key))
            ws.cell(row=row, column=2, value=str(_val(node) or ""))
            ws.cell(row=row, column=3, value=_conf(node))
            row += 1

        # Medications table
        row += 1
        meds = result.get("medications") or []
        if meds:
            ws.cell(row=row, column=1, value="الأدوية / Medications").font = hdr_font
            row += 1
            cols = ["name", "drug_class", "form", "dosage", "duration", "confidence"]
            for c, col in enumerate(cols, start=1):
                cell = ws.cell(row=row, column=c, value=col)
                cell.font = hdr_font
                cell.fill = hdr_fill
            row += 1
            for m in meds:
                for c, col in enumerate(cols, start=1):
                    v = m.get(col)
                    if col == "confidence":
                        v = _conf_str(v)
                    ws.cell(row=row, column=c, value=str(v) if v is not None else "")
                row += 1

        # As-needed
        prn = result.get("as_needed") or []
        if prn:
            row += 1
            ws.cell(row=row, column=1, value="حسب الحاجة / As-needed").font = hdr_font
            row += 1
            for p in prn:
                ws.cell(row=row, column=1, value=str(p.get("name", "")))
                ws.cell(row=row, column=2, value=_conf_str(p.get("confidence")))
                row += 1

        # Validations
        vals = result.get("validations") or []
        if vals:
            row += 1
            ws.cell(row=row, column=1, value="فحوصات / Validations").font = hdr_font
            row += 1
            for v in vals:
                mark = "✓" if v.get("ok") else "⚠"
                ws.cell(row=row, column=1, value=mark)
                ws.cell(row=row, column=2, value=str(v.get("check", "")))
                ws.cell(row=row, column=3, value=str(v.get("detail", "")))
                row += 1

        for col_letter in ("A", "B", "C", "D", "E", "F"):
            ws.column_dimensions[col_letter].width = 22

        buf = io.BytesIO(); wb.save(buf)
        return (buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                f"{title}.xlsx")

    if layout == "mixed":
        # Header + line items + totals (invoice / receipt shape)
        wb = Workbook(); ws = wb.active
        ws.title = "Invoice"
        ws.sheet_view.rightToLeft = True
        hdr_font = Font(bold=True)
        hdr_fill = PatternFill("solid", fgColor="DBEAFE")

        ws["A1"] = doc_type.replace("_", " ").title()
        ws["A1"].font = Font(bold=True, size=13)

        row = 3
        header = result.get("header") or {}
        if header:
            ws.cell(row=row, column=1, value="Header").font = hdr_font
            row += 1
            for k, node in header.items():
                ws.cell(row=row, column=1, value=str(k))
                ws.cell(row=row, column=2, value=str(_val(node) or ""))
                ws.cell(row=row, column=3, value=_conf(node))
                row += 1

        items = result.get("line_items") or {}
        cols = items.get("columns") or []
        rows_ = items.get("rows") or []
        if cols and rows_:
            row += 1
            ws.cell(row=row, column=1, value="Line items").font = hdr_font
            row += 1
            for c, col in enumerate(cols, start=1):
                cell = ws.cell(row=row, column=c, value=str(col))
                cell.font = hdr_font; cell.fill = hdr_fill
            row += 1
            for r_ in rows_:
                for c, val in enumerate(r_, start=1):
                    ws.cell(row=row, column=c, value=("" if val is None else str(val)))
                row += 1

        totals = result.get("totals") or {}
        if totals:
            row += 1
            ws.cell(row=row, column=1, value="Totals").font = hdr_font
            row += 1
            for k, node in totals.items():
                ws.cell(row=row, column=1, value=str(k))
                ws.cell(row=row, column=2, value=str(_val(node) or ""))
                ws.cell(row=row, column=3, value=_conf(node))
                row += 1

        for col_letter in ("A", "B", "C", "D", "E", "F"):
            ws.column_dimensions[col_letter].width = 22

        buf = io.BytesIO(); wb.save(buf)
        return (buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                f"{title}.xlsx")

    if layout == "table":
        tbl = result.get("table") or {}
        return table_to_xlsx(tbl.get("columns") or [], tbl.get("rows") or [], title)

    # Unknown layout — flatten fields and use generic exporter.
    fields = result.get("fields") if isinstance(result, dict) else None
    rows = []
    if isinstance(fields, dict):
        for k, node in fields.items():
            rows.append({
                "page": "",
                "field": str(k),
                "value": _val(node),
                "confidence": (node or {}).get("confidence") if isinstance(node, dict) else None,
            })
    return to_xlsx(rows, title)


# ---- Table exports (columns + rows) ----

def table_to_xlsx(columns, rows, title="table"):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.sheet_view.rightToLeft = True
    if columns:
        ws.append([str(c) for c in columns])
    for r in rows:
        ws.append([("" if c is None else str(c)) for c in r])
    buf = io.BytesIO()
    wb.save(buf)
    return (buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            f"{title}.xlsx")


def table_to_csv(columns, rows, title="table"):
    buf = io.StringIO()
    w = csv.writer(buf)
    if columns:
        w.writerow(columns)
    for r in rows:
        w.writerow(r)
    return ("﻿" + buf.getvalue()).encode("utf-8"), "text/csv; charset=utf-8", f"{title}.csv"


# ---- Plain text -> Word / PDF ----

def text_to_docx(text, title="document"):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for line in str(text).split("\n"):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    buf = io.BytesIO()
    doc.save(buf)
    return (buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{title}.docx")


def text_to_pdf(text, title="document"):
    global _FONT_READY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    if not _FONT_READY:
        pdfmetrics.registerFont(TTFont("Naskh", ARABIC_FONT))
        _FONT_READY = True
    style = ParagraphStyle("p", fontName="Naskh", fontSize=12, alignment=2, wordWrap="RTL", leading=18)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=title)
    flow = []
    for line in str(text).split("\n"):
        flow.append(Paragraph(_ar(line) if line.strip() else "&nbsp;", style))
    doc.build(flow or [Spacer(1, 12)])
    return buf.getvalue(), "application/pdf", f"{title}.pdf"


# ---- Searchable PDF (image + invisible text layer) ----

def image_to_searchable_pdf(image_bytes, text, title="searchable"):
    global _FONT_READY
    from PIL import Image
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    if not _FONT_READY:
        pdfmetrics.registerFont(TTFont("Naskh", ARABIC_FONT))
        _FONT_READY = True

    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    w, h = img.size
    png = io.BytesIO()
    img.save(png, format="PNG")
    png.seek(0)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(w, h))
    c.drawImage(ImageReader(png), 0, 0, width=w, height=h)
    to = c.beginText(4, h - 14)
    to.setFont("Naskh", 10)
    to.setTextRenderMode(3)  # invisible but selectable/searchable
    for line in str(text).split("\n"):
        to.textLine(_ar(line))
    c.drawText(to)
    c.showPage()
    c.save()
    return buf.getvalue(), "application/pdf", f"{title}.pdf"
