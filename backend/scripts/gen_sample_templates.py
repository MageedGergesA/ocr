"""Generate 5 sample templates (xlsx/docx ×2/html/pdf) for the upload-template
demo. Re-run when the schema or styling changes:
    python backend/scripts/gen_sample_templates.py
Output goes to backend/app/static/sample-templates/ which is served by the
StaticFiles mount and linked from /app's gallery."""
import io
import os
from pathlib import Path

OUT = str(Path(__file__).resolve().parents[1] / "app/static/sample-templates")
os.makedirs(OUT, exist_ok=True)

# ---- 1. Invoice template (xlsx) ----
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

wb = Workbook()
ws = wb.active
ws.title = "Invoice"
headers = ["Vendor Name", "Invoice Number", "Issue Date", "Due Date",
           "Currency", "Subtotal", "Tax Amount", "Total"]
ws.append(headers)
# Header styling
header_font = Font(bold=True, color="FFFFFF", size=11)
header_fill = PatternFill("solid", fgColor="1B5E63")
thin = Side(border_style="thin", color="888888")
border = Border(left=thin, right=thin, top=thin, bottom=thin)
for col_idx in range(1, len(headers) + 1):
    cell = ws.cell(row=1, column=col_idx)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center", vertical="center")
    cell.border = border
    ws.column_dimensions[cell.column_letter].width = 18
ws.row_dimensions[1].height = 28
# Add a sample second sheet with line items
ws2 = wb.create_sheet("Line Items")
ws2.append(["Description", "Quantity", "Unit Price", "Total"])
for col_idx in range(1, 5):
    cell = ws2.cell(row=1, column=col_idx)
    cell.font = header_font; cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center"); cell.border = border
    ws2.column_dimensions[cell.column_letter].width = 22
wb.save(os.path.join(OUT, "invoice-template.xlsx"))
print("OK invoice-template.xlsx")

# ---- 2. Patient intake form (docx) ----
# Use a row-1 = headers structure so fill_docx writes row-2 values under them.
# Split into 2 horizontal tables so the page stays narrow enough.
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()
title = doc.add_paragraph()
title_run = title.add_run("Patient Intake Form")
title_run.font.size = Pt(20); title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1b, 0x5e, 0x63)
doc.add_paragraph("Single-record template — extraction fills row 2 under each header column.")
doc.add_paragraph()
fields = [
    "Patient Name", "Date of Birth", "National ID", "Phone Number",
    "Email", "Address", "Insurance Provider", "Insurance Number",
    "Emergency Contact", "Blood Group", "Allergies", "Chronic Conditions",
    "Current Medications", "Chief Complaint", "Blood Pressure", "Visit Date",
]
half = len(fields) // 2
for chunk in (fields[:half], fields[half:]):
    tbl = doc.add_table(rows=2, cols=len(chunk))
    tbl.style = "Light Grid Accent 1"
    for col_idx, name in enumerate(chunk):
        cell = tbl.rows[0].cells[col_idx]
        cell.text = name
        for run in cell.paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()
doc.add_paragraph("Doctor: ___________   Date: ___________   Signature: ___________")
doc.save(os.path.join(OUT, "patient-intake.docx"))
print("OK patient-intake.docx")

# ---- 3. Real estate contract template (docx with placeholders) ----
doc = Document()
h = doc.add_paragraph()
hr = h.add_run("Real Estate Sale Contract")
hr.font.size = Pt(18); hr.font.bold = True
hr.font.color.rgb = RGBColor(0x1b, 0x5e, 0x63)
doc.add_paragraph()

# Use a Field|Value table so the docx parser detects fields
tbl = doc.add_table(rows=8, cols=2)
tbl.style = "Light Grid Accent 1"
fields = ["Seller Name", "Buyer Name", "Property Address", "Unit Number",
          "Property Area", "Sale Price", "Payment Method", "Contract Date"]
tbl.rows[0].cells[0].text = "Field"; tbl.rows[0].cells[1].text = "Value"
for r in tbl.rows[0].cells[0].paragraphs[0].runs + tbl.rows[0].cells[1].paragraphs[0].runs:
    r.bold = True
for i, name in enumerate(fields[1:], start=1):
    tbl.rows[i].cells[0].text = name
# Put first field at row 1 too (one record)
tbl.rows[0].cells[0].text = fields[0]  # Override header to be first field

# Rebuild as a single-row record table so fill_docx works
doc = Document()
h = doc.add_paragraph()
hr = h.add_run("Real Estate Sale Contract")
hr.font.size = Pt(18); hr.font.bold = True
hr.font.color.rgb = RGBColor(0x1b, 0x5e, 0x63)
doc.add_paragraph("Single-record template — extraction fills row 2 under each column header.")
doc.add_paragraph()
record_tbl = doc.add_table(rows=2, cols=len(fields))
record_tbl.style = "Light Grid Accent 1"
for col_idx, name in enumerate(fields):
    cell = record_tbl.rows[0].cells[col_idx]
    cell.text = name
    for run in cell.paragraphs[0].runs:
        run.bold = True

doc.add_paragraph()
doc.add_paragraph("Notes paragraph with {{ seller_name }} and price of {{ sale_price }} EGP "
                  "between {{ buyer_name }} and seller; signed on {{ contract_date }}.")
doc.add_paragraph()
doc.add_paragraph("Witness 1: ___________________   Witness 2: ___________________")

doc.save(os.path.join(OUT, "sale-contract.docx"))
print("OK sale-contract.docx")

# ---- 4. Patient registration form (HTML) ----
html = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Patient Registration Form</title>
<style>
body { font-family: 'Segoe UI', sans-serif; max-width: 720px; margin: 40px auto; padding: 0 20px; color: #1f2937; }
h1 { color: #1B5E63; border-bottom: 2px solid #1B5E63; padding-bottom: 8px; }
.row { display: grid; grid-template-columns: 200px 1fr; gap: 12px; margin: 10px 0; align-items: center; }
label { font-weight: 600; color: #374151; }
input, select { padding: 8px 10px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 14px; }
table { width: 100%; border-collapse: collapse; margin: 18px 0; }
th, td { border: 1px solid #d1d5db; padding: 8px; text-align: left; }
th { background: #f3f4f6; color: #1B5E63; }
</style>
</head>
<body>
<h1>Patient Registration Form</h1>
<p>Please complete all fields. Upload this form to Mostakhles with the patient's ID + insurance card to auto-fill.</p>
<form>
  <div class="row"><label for="full_name">Full Name</label><input id="full_name" name="full_name" type="text" placeholder="e.g. Mohamed Ali Hassan"></div>
  <div class="row"><label for="date_of_birth">Date of Birth</label><input id="date_of_birth" name="date_of_birth" type="date"></div>
  <div class="row"><label for="national_id">National ID</label><input id="national_id" name="national_id" type="text" placeholder="14 digits"></div>
  <div class="row"><label for="phone_number">Phone Number</label><input id="phone_number" name="phone_number" type="tel" placeholder="+20 1XX XXX XXXX"></div>
  <div class="row"><label for="email">Email</label><input id="email" name="email" type="email"></div>
  <div class="row"><label for="address">Address</label><input id="address" name="address" type="text"></div>
  <div class="row"><label for="insurance_provider">Insurance Provider</label><input id="insurance_provider" name="insurance_provider" type="text"></div>
  <div class="row"><label for="insurance_number">Insurance Number</label><input id="insurance_number" name="insurance_number" type="text"></div>
  <div class="row"><label for="emergency_contact">Emergency Contact</label><input id="emergency_contact" name="emergency_contact" type="text"></div>
  <div class="row"><label for="blood_group">Blood Group</label><input id="blood_group" name="blood_group" type="text" placeholder="A+, O-, ..."></div>
  <div class="row"><label for="allergies">Allergies</label><input id="allergies" name="allergies" type="text"></div>
  <div class="row"><label for="chronic_conditions">Chronic Conditions</label><input id="chronic_conditions" name="chronic_conditions" type="text"></div>
</form>
<h3>Medication History</h3>
<table>
  <thead><tr><th>Medication Name</th><th>Dosage</th><th>Frequency</th><th>Start Date</th></tr></thead>
  <tbody><tr><td></td><td></td><td></td><td></td></tr></tbody>
</table>
</body>
</html>
"""
with open(os.path.join(OUT, "patient-registration.html"), "w", encoding="utf-8") as f:
    f.write(html)
print("OK patient-registration.html")

# ---- 5. Egyptian tax declaration form (PDF) ----
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

pdf_path = os.path.join(OUT, "tax-declaration.pdf")
c = canvas.Canvas(pdf_path, pagesize=A4)
W, H = A4
# Header band
c.setFillColorRGB(0.106, 0.369, 0.388)  # brand teal
c.rect(0, H - 3*cm, W, 3*cm, fill=1, stroke=0)
c.setFillColorRGB(1, 1, 1)
c.setFont("Helvetica-Bold", 20)
c.drawString(2*cm, H - 1.7*cm, "Tax Declaration Form")
c.setFont("Helvetica", 11)
c.drawString(2*cm, H - 2.5*cm, "Egyptian Tax Authority - Annual Filing")

# Form fields (label: ___ pattern that our pdfplumber fallback scrapes)
c.setFillColorRGB(0.1, 0.1, 0.1)
y = H - 5*cm
fields = [
    ("Taxpayer Name", "Full legal or company name"),
    ("Tax Registration Number", "TRN issued by ETA"),
    ("Commercial Register Number", "CR number if applicable"),
    ("National ID", "14-digit individual ID"),
    ("Activity Type", "Main commercial activity"),
    ("Tax Period", "Fiscal year (e.g. 2025)"),
    ("Total Revenue", "Annual gross revenue (EGP)"),
    ("Deductible Expenses", "Total deductions claimed (EGP)"),
    ("Taxable Income", "Net taxable income (EGP)"),
    ("Tax Due", "Calculated tax owed (EGP)"),
    ("Filing Date", "DD/MM/YYYY"),
    ("Submitted By", "Name of person filing"),
]
for label, hint in fields:
    c.setFont("Helvetica-Bold", 10)
    c.drawString(2*cm, y, f"{label}:")
    c.setFont("Helvetica-Oblique", 8)
    c.setFillColorRGB(0.5, 0.5, 0.5)
    c.drawString(2*cm + 5.5*cm, y, hint)
    c.setFillColorRGB(0.1, 0.1, 0.1)
    # Field line
    c.setStrokeColorRGB(0.7, 0.7, 0.7)
    c.line(2*cm + 12*cm, y, W - 2*cm, y)
    y -= 1.0*cm

# Signature box at bottom
y -= 1*cm
c.setFont("Helvetica-Bold", 10)
c.drawString(2*cm, y, "Signature:")
c.line(2*cm + 3*cm, y, 2*cm + 10*cm, y)
c.drawString(11*cm, y, "Stamp:")
c.rect(13*cm, y - 1.5*cm, 4*cm, 2*cm, fill=0)

c.showPage()
c.save()
print("OK tax-declaration.pdf")

print("\nAll 5 sample templates generated.")
